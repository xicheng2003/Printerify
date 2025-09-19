# printerify/api/services/pricing.py

import os
import sys
import math
import shutil
import zipfile
import tempfile
import subprocess
import xml.etree.ElementTree as ET

import fitz  # PyMuPDF, 用于读取PDF
import docx  # python-docx, 用于读取DOCX
from decimal import Decimal # 使用 Decimal 来精确计算货币

# --- 价格配置 (已重构) ---
# 价格结构现在以“纸张尺寸”作为顶层key，更具扩展性
PRICE_CONFIG = {
    'base_service_fee': Decimal('0.50'),
    'print': {
        'a4_70g': {
            'black_white': {
                'single': Decimal('0.15'),
                'double': Decimal('0.15'),
                'single_double': None, # 沿用单/双面价格逻辑
            },
            'color': {
                'single': Decimal('0.50'),
                'double': Decimal('0.80'),
                'single_double': None, # 封面单面+内容双面
            }
        },
        'a4_80g': {
            'black_white': {
                'single': Decimal('0.15'),
                'double': Decimal('0.15'),
                'single_double': None,
            },
            'color': {
                'single': Decimal('0.50'),
                'double': Decimal('0.80'),
                'single_double': None,
            }
        },
        'b5_70g': {
            'black_white': {
                'single': Decimal('0.12'), # B5 价格稍低
                'double': Decimal('0.12'),
                'single_double': None,
            },
            'color': {
                'single': Decimal('0.40'),
                'double': Decimal('0.70'),
                'single_double': None,
            }
        }
        # 您未来可以在此添加 'a3': { ... } 等更多规格
    },
    'binding': {
        'none': Decimal('0.00'),
        'staple_top_left': Decimal('0.10'),
        'staple_left_side': Decimal('0.10'),
        'staple': Decimal('2.00'),
        'ring_bound': Decimal('5.00'),
    }
}

# --- 核心计费函数 ---

def _calculate_pages(file_path, prefer_exact: bool = False):
    """
    计算文件页数的核心逻辑。
    【已延续】此函数延续了您项目中对不同文件类型的处理方式。
    """
    try:
        if file_path.lower().endswith('.pdf'):
            with fitz.open(file_path) as doc:
                return doc.page_count
        elif file_path.lower().endswith('.docx'):
            return _docx_page_count(file_path, prefer_exact=prefer_exact)
        elif file_path.lower().endswith('.doc'):
            return _doc_page_count(file_path, prefer_exact=prefer_exact)
        else:
            # 对于不支持的格式，暂时返回1页，避免程序崩溃
            return 1
    except Exception as e:
        print(f"Error calculating pages for {file_path}: {e}")
        # 如果计算出错，返回一个默认值，并打印错误
        return 1


# --- DOCX 页数计算（多策略，按可靠性降级） ---
def _docx_page_count(file_path: str, prefer_exact: bool = False) -> int:
    """
    更可靠的 DOCX 页数计算：
    1) 尝试读取 docx 内置元数据 docProps/app.xml 的 Pages 字段（最快、无依赖）
    2) 尝试使用 LibreOffice/soffice 无头转换为 PDF 后读取页数（跨平台，推荐在服务器上安装）
    3) 在 Windows 且安装了 Word 的情况下，使用 COM 的 ComputeStatistics 获取精确页数
    4) 可选：使用 docx2pdf 转为 PDF 后由 PyMuPDF 读取页数（通常依赖本机 Word）
    4) 最后回退：基于文本、表格与图片的启发式估算

    始终返回 >= 1 的整数。
    """
    # 1) 读取 docProps/app.xml
    pages = _docx_pages_from_app_xml(file_path)
    if isinstance(pages, int) and pages > 0:
        return pages

    # 2) LibreOffice/soffice 转 PDF 后计数（跨平台，可通过环境/配置开关启用）
    if prefer_exact or _is_soffice_enabled():
        pages = _docx_pages_via_soffice(file_path, timeout_sec=_get_soffice_timeout())
        if isinstance(pages, int) and pages > 0:
            return pages

    # 3) Windows + Word COM（pywin32）
    pages = _docx_pages_from_word_com(file_path)
    if isinstance(pages, int) and pages > 0:
        return pages

    # 4) docx2pdf 转 PDF 再数页（可选，可通过环境/配置开关启用）
    if _is_docx2pdf_enabled():
        pages = _docx_pages_via_pdf(file_path)
        if isinstance(pages, int) and pages > 0:
            return pages

    # 5) 启发式估算
    pages = _docx_pages_heuristic(file_path)
    return pages if pages and pages > 0 else 1


def _doc_page_count(file_path: str, prefer_exact: bool = False) -> int:
    """
    兼容旧版 .doc（二进制）文件的页数计算。
    优先策略：
    1) LibreOffice/soffice 转 PDF 读取页数（跨平台，推荐安装）
    2) Windows + Word COM（pywin32）ComputeStatistics（若可用）
    3) 可选：docx2pdf（通常依赖 Word）再读取页数
    4) 最后回退：返回 1（无法可靠解析 .doc 内容结构）

    始终返回 >= 1 的整数。
    """
    # 1) LibreOffice/soffice（支持 .doc）
    if prefer_exact or _is_soffice_enabled():
        pages = _docx_pages_via_soffice(file_path, timeout_sec=_get_soffice_timeout())
        if isinstance(pages, int) and pages > 0:
            return pages

    # 2) Windows + Word COM（pywin32）
    pages = _docx_pages_from_word_com(file_path)
    if isinstance(pages, int) and pages > 0:
        return pages

    # 3) docx2pdf（Word 驱动，可能支持 .doc，若不可用将返回 None）
    if _is_docx2pdf_enabled():
        pages = _docx_pages_via_pdf(file_path)
        if isinstance(pages, int) and pages > 0:
            return pages

    # 4) 回退：无法可靠估算 .doc，返回 1 以避免阻塞
    return 1


def _docx_pages_from_app_xml(file_path: str) -> int | None:
    """从 docx 的 docProps/app.xml 读取 Pages 字段（如果存在）。"""
    try:
        with zipfile.ZipFile(file_path) as zf:
            with zf.open('docProps/app.xml') as f:
                data = f.read()
        root = ET.fromstring(data)
        # 兼容命名空间，匹配以 'Pages' 结尾的标签
        for el in root.iter():
            if el.tag.endswith('Pages') and el.text:
                try:
                    val = int(el.text.strip())
                    if val > 0:
                        return val
                except ValueError:
                    pass
    except Exception:
        pass
    return None


def _docx_pages_from_word_com(file_path: str) -> int | None:
    """
    使用 Windows Word COM 获取精确页数。
    要求：
    - 运行在 Windows 上（sys.platform.startswith('win')）
    - 已安装 Microsoft Word
    - 可选依赖 pywin32（win32com）
    """
    if not sys.platform.startswith('win'):
        return None

    word = None
    doc = None
    try:
        try:
            import win32com.client  # type: ignore
            from win32com.client import constants  # type: ignore
        except Exception:
            return None

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        # 为避免弹窗，设置 DisplayAlerts
        try:
            word.DisplayAlerts = 0  # wdAlertsNone
        except Exception:
            pass

        abs_path = os.path.abspath(file_path)
        doc = word.Documents.Open(abs_path, ReadOnly=True)
        # 强制重新分页，获取最新统计
        try:
            doc.Repaginate()
        except Exception:
            pass

        pages = int(doc.ComputeStatistics(constants.wdStatisticPages, True))
        return pages if pages > 0 else None
    except Exception:
        return None
    finally:
        try:
            if doc is not None:
                doc.Close(False)
        except Exception:
            pass
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass


def _docx_pages_via_soffice(file_path: str, timeout_sec: int = 60) -> int | None:
    """
    使用 LibreOffice/soffice 无头模式将 DOCX 转为 PDF，再读取页数。
    需要服务器安装 libreoffice（可执行为 `soffice` 或 `libreoffice`）。
    若命令不可用或转换失败，则返回 None。
    """
    candidates = ['soffice', 'libreoffice']
    tmpdir = None
    try:
        # 查找可用命令
        cmd = next((c for c in candidates if shutil.which(c)), None)
        if not cmd:
            return None

        tmpdir = tempfile.mkdtemp(prefix='soffice_')
        # LibreOffice 会在 outdir 生成同名 .pdf
        outdir = tmpdir
        result = subprocess.run(
            [cmd, '--headless', '--norestore', '--convert-to', 'pdf', '--outdir', outdir, os.path.abspath(file_path)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout_sec,
            check=False,
        )
        if result.returncode != 0:
            # 打印一次错误方便诊断，但不中断总体流程
            try:
                err = result.stderr.decode(errors='ignore')
                print(f"LibreOffice conversion error: {err}")
            except Exception:
                pass
            return None

        base = os.path.splitext(os.path.basename(file_path))[0]
        pdf_path = os.path.join(outdir, base + '.pdf')
        if not os.path.exists(pdf_path):
            return None

        with fitz.open(pdf_path) as pdf:
            pages = int(pdf.page_count)
            return pages if pages > 0 else None
    except subprocess.TimeoutExpired:
        print(f"LibreOffice conversion timeout for {file_path}")
        return None
    except Exception:
        return None
    finally:
        if tmpdir and os.path.isdir(tmpdir):
            try:
                shutil.rmtree(tmpdir)
            except Exception:
                pass


def _is_soffice_enabled() -> bool:
    """是否启用 LibreOffice 转换（默认关闭，以保证同步接口快速返回）。
    优先读取环境变量 DOCX_SOFFICE_ENABLED，其次读取 Django settings.DOCX_PAGECOUNT_USE_SOFFICE。
    """
    env_val = os.getenv('DOCX_SOFFICE_ENABLED', '').strip().lower()
    if env_val in {'1', 'true', 'yes', 'on'}:
        return True
    if env_val in {'0', 'false', 'no', 'off', ''}:
        # 留空时再看 settings
        pass
    try:
        from django.conf import settings  # type: ignore
        return bool(getattr(settings, 'DOCX_PAGECOUNT_USE_SOFFICE', False))
    except Exception:
        return False


def _get_soffice_timeout() -> int:
    """获取 LibreOffice 转换超时（秒）。默认 10s，避免阻塞同步请求太久。"""
    val = os.getenv('DOCX_SOFFICE_TIMEOUT', '').strip()
    if val.isdigit():
        return max(1, int(val))
    try:
        from django.conf import settings  # type: ignore
        return int(getattr(settings, 'DOCX_PAGECOUNT_SOFFICE_TIMEOUT', 10))
    except Exception:
        return 10


def _is_docx2pdf_enabled() -> bool:
    """是否启用 docx2pdf 转换（默认关闭，通常依赖 Word，不适合服务器）。
    读取环境变量 DOCX_DOCX2PDF_ENABLED 或 settings.DOCX_PAGECOUNT_USE_DOCX2PDF。
    """
    env_val = os.getenv('DOCX_DOCX2PDF_ENABLED', '').strip().lower()
    if env_val in {'1', 'true', 'yes', 'on'}:
        return True
    if env_val in {'0', 'false', 'no', 'off', ''}:
        pass
    try:
        from django.conf import settings  # type: ignore
        return bool(getattr(settings, 'DOCX_PAGECOUNT_USE_DOCX2PDF', False))
    except Exception:
        return False


def _docx_pages_via_pdf(file_path: str) -> int | None:
    """
    使用 docx2pdf 将 DOCX 转为 PDF，再用 PyMuPDF 读取页数。
    注意：在 Windows 上通常仍需本机安装 Word。
    """
    try:
        from docx2pdf import convert  # type: ignore
    except Exception:
        return None

    tmpdir = None
    try:
        tmpdir = tempfile.mkdtemp(prefix='docx2pdf_')
        out_pdf = os.path.join(tmpdir, 'tmp.pdf')
        convert(os.path.abspath(file_path), out_pdf)
        with fitz.open(out_pdf) as pdf:
            pages = int(pdf.page_count)
        return pages if pages > 0 else None
    except Exception:
        return None
    finally:
        if tmpdir and os.path.isdir(tmpdir):
            try:
                shutil.rmtree(tmpdir)
            except Exception:
                pass


def _docx_pages_heuristic(file_path: str) -> int:
    """
    最后兜底的启发式估算：
    - 基于段落文本的词数估算（~450 词/页）
    - 图片（inline shapes）每 4 张算 1 页
    - 表格单元格数量每 30 个算 1 页
    该估算不精确，但通常比“sections 数量”更合理。
    """
    try:
        d = docx.Document(file_path)
        word_count = 0
        for p in d.paragraphs:
            t = (p.text or '').strip()
            if t:
                word_count += len(t.split())

        # 图片数量（尽力而为，某些场景可能不可用）
        images = 0
        try:
            images = len(d.inline_shapes)
        except Exception:
            pass

        # 表格单元格数量
        cells = 0
        try:
            for table in d.tables:
                for row in table.rows:
                    cells += len(row.cells)
        except Exception:
            pass

        pages_text = math.ceil(word_count / 450) if word_count > 0 else 0
        pages_images = math.ceil(images / 4) if images > 0 else 0
        pages_tables = math.ceil(cells / 30) if cells > 0 else 0

        estimated = pages_text + pages_images + pages_tables
        return max(1, estimated)
    except Exception:
        return 1

# --- ▼▼▼ calculate_document_price 函数已重构 ▼▼▼ ---
def calculate_document_price(file_path, paper_size, color_mode, print_sided, copies, *, prefer_exact: bool = False, override_page_count: int | None = None):
    """
    计算【单个文件】的打印费用 (已更新以支持纸张尺寸和新的单双面选项)。
    返回: (页数, 打印费用) 的元组
    """
    page_count = int(override_page_count) if override_page_count else _calculate_pages(file_path, prefer_exact=prefer_exact)
    if page_count == 0:
        return 0, Decimal('0.00')

    total_print_cost = Decimal('0.00')

    try:
        # 1. 根据纸张尺寸和颜色获取价格表
        try:
            price_map = PRICE_CONFIG['print'][paper_size][color_mode]
        except KeyError:
            # 兼容裸纸型（例如 'a4'、'b5'）的情况，映射到默认 70g
            fallback_map = {
                'a4': 'a4_70g',
                'b5': 'b5_70g',
            }
            mapped = fallback_map.get(paper_size, paper_size)
            price_map = PRICE_CONFIG['print'][mapped][color_mode]
    except KeyError:
        # 如果配置不存在 (例如一个无效的 paper_size)，返回0价
        return page_count, Decimal('0.0')

    # 2. 根据打印方式计算价格
    if print_sided == 'single_double':
        # “封面单面”的特殊逻辑
        if page_count == 1:
            # 只有一页时，按单面计算
            total_print_cost = price_map['single']
        else:
            # 封面按“单面”价，其余 (page_count - 1) 页按“双面”价
            cover_cost = price_map['single']
            content_cost = (page_count - 1) * price_map['double']
            total_print_cost = cover_cost + content_cost
    else:
        # 标准的单面或双面逻辑
        price_per_page = price_map.get(print_sided, Decimal('0.00'))
        total_print_cost = price_per_page * page_count

    # 3. 乘以份数得到最终价格
    final_cost = total_print_cost * copies
    
    return page_count, final_cost
# --- ▲▲▲ 函数重构结束 ▲▲▲ ---

def calculate_binding_cost(binding_type, page_count):
    """
    计算【单个装订组】的装订费用。
    这是被 Serializer 调用的主要函数之二。
    """
    # 简单的实现：如果页数大于0且需要装订，则收取装订费
    if page_count > 0 and binding_type != 'none':
        try:
            return PRICE_CONFIG['binding'][binding_type]
        except KeyError:
            return Decimal('0.00') # 如果装订类型不存在，则不收费
    
    # 如果不满足装订条件，则装订费为0
    return Decimal('0.00')